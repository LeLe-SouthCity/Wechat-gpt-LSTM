
import json
from docx import Document as DocxDocument



def convert_json_to_word(json_file_path:str, word_file_path:str):
    """
    json_file_path(str):json文件路径
    word_file_path(str):保存的word文件路径
    """
    with open(json_file_path, 'r', encoding='utf-8') as json_file:
        data = json.load(json_file)

    doc = DocxDocument()

    try:
        if isinstance(data, dict):
            for key, value in data.items():
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"{key}: ").bold = True
                paragraph.add_run(str(value))
        elif isinstance(data, list):
            for item in data:
                doc.add_paragraph(str(item))
        else:
            doc.add_paragraph(str(data))
    except UnicodeEncodeError as e:
        print(f"Encoding error: {e}")
        # 处理编码错误或替换问题字符的代码

    doc.save(word_file_path)
 